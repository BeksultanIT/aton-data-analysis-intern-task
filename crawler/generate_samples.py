import os
import zipfile
import io

BASE_DIR = "storage"
os.makedirs(BASE_DIR, exist_ok=True)

def make_txt():
    content = """Заметки аналитика — Q1 2026

Итоги первого квартала показывают устойчивый рост по сегменту инвестиционных продуктов.
Клиенты с высоким капиталом обеспечили около 63% общей выручки.

Ключевые метрики:
- Средняя сумма транзакции: 245 000 руб.
- Количество активных клиентов: 3 841
- Самая популярная услуга: Доверительное управление

Рекомендации на Q2: фокус на привлечение новых клиентов в сегмент "Средний капитал".
"""
    path = os.path.join(BASE_DIR, "notes.txt")
    with open(path, "w", encoding="utf-8") as f:
        f.write(content)
    print(f"Создан: {path}")
    return path


def make_docx():
    try:
        import docx
    except ImportError:
        print("python-docx не установлен, .docx пропускаем")
        return None

    doc = docx.Document()
    doc.add_heading("Квартальный отчёт АТОН — Q1 2026", level=1)
    doc.add_paragraph(
        "Настоящий отчёт подготовлен группой анализа данных на основе "
        "транзакционных данных за январь–март 2026 года."
    )
    doc.add_heading("Выручка по сервисам", level=2)
    doc.add_paragraph(
        "Лидером по выручке стала услуга «Структурирование капитала» — "
        "её доля составила 28% от общего объёма поступлений."
    )
    doc.add_heading("Клиентская база", level=2)
    doc.add_paragraph(
        "Всего за отчётный период зафиксировано 10 001 транзакция "
        "от 8 327 уникальных клиентов. Доля клиентов с высоким капиталом "
        "(более 1 млн руб. чистых активов) увеличилась на 4 п.п. по сравнению с Q4 2025."
    )

    path = os.path.join(BASE_DIR, "report_q1.docx")
    doc.save(path)
    print(f"Создан: {path}")
    return path


def make_xlsx():
    try:
        import openpyxl
    except ImportError:
        print("openpyxl не установлен, .xlsx пропускаем")
        return None

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Клиенты"

    headers = ["client_id", "age", "gender", "net_worth", "asset_level"]
    ws.append(headers)

    sample_data = [
        (1001, 34, "М", 450_000, "Средний капитал"),
        (1002, 52, "Ж", 2_100_000, "Высокий капитал"),
        (1003, 28, "М", 75_000, "Низкий капитал"),
        (1004, 45, "Ж", 850_000, "Средний капитал"),
        (1005, 61, "М", 5_400_000, "Высокий капитал"),
    ]
    for row in sample_data:
        ws.append(list(row))

    ws2 = wb.create_sheet("Транзакции")
    ws2.append(["transaction_id", "client_id", "service", "amount", "city"])
    ws2.append([10001, 1001, "Доверительное управление", 120_000, "Москва"])
    ws2.append([10002, 1002, "Структурирование капитала", 500_000, "Санкт-Петербург"])
    ws2.append([10003, 1003, "Брокерское обслуживание", 15_000, "Казань"])

    path = os.path.join(BASE_DIR, "data_clients.xlsx")
    wb.save(path)
    print(f"Создан: {path}")
    return path



def make_pdf():
    try:
        from reportlab.pdfgen import canvas
        from reportlab.lib.pagesizes import A4
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont

        path = os.path.join(BASE_DIR, "summary.pdf")
        c = canvas.Canvas(path, pagesize=A4)
        width, height = A4

        c.setFont("Helvetica-Bold", 16)
        c.drawString(60, height - 80, "ATON — Investment Summary Q1 2026")

        c.setFont("Helvetica", 12)
        lines = [
            "Total revenue: 124,500,000 RUB",
            "Active clients: 8 327",
            "Top service by volume: Capital Structuring",
            "Top city by avg. transaction: Moscow (310 000 RUB avg.)",
            "",
            "Forecast for next month: 11 200 000 RUB (linear regression model)",
        ]
        y = height - 130
        for line in lines:
            c.drawString(60, y, line)
            y -= 22

        c.save()
        print(f"Создан: {path}")
        return path

    except ImportError:
        path = os.path.join(BASE_DIR, "summary_placeholder.txt")
        with open(path, "w", encoding="utf-8") as f:
            f.write(
                "PDF-файл: ATON Investment Summary Q1 2026\n"
                "Total revenue: 124,500,000 RUB\n"
                "Active clients: 8 327\n"
                "(reportlab не установлен, файл создан как txt-заглушка)\n"
            )
        print(f"Создан (txt-заглушка вместо pdf): {path}")
        return path



def make_zip(inner_files: list[str]):
    """Упаковываем несколько файлов в zip-архив."""
    path = os.path.join(BASE_DIR, "archive.zip")
    with zipfile.ZipFile(path, "w", zipfile.ZIP_DEFLATED) as zf:
        for fp in inner_files:
            if fp and os.path.exists(fp):
                zf.write(fp, arcname=os.path.basename(fp))

    print(f"Создан архив: {path} (файлов: {len(inner_files)})")
    return path



def make_7z(inner_files: list[str]):
    try:
        import py7zr
    except ImportError:
        print("py7zr не установлен, .7z пропускаем")
        return None

    path = os.path.join(BASE_DIR, "archive.7z")
    with py7zr.SevenZipFile(path, mode="w") as sz:
        for fp in inner_files:
            if fp and os.path.exists(fp):
                sz.write(fp, arcname=os.path.basename(fp))

    print(f"Создан архив: {path}")
    return path



def main():
    print(f"Генерируем тестовое хранилище в директории '{BASE_DIR}/'...\n")

    txt_path  = make_txt()
    docx_path = make_docx()
    xlsx_path = make_xlsx()
    pdf_path  = make_pdf()

    files_for_archive = [f for f in [txt_path, docx_path, xlsx_path] if f]
    make_zip(files_for_archive)
    make_7z([pdf_path] if pdf_path else [])

    print(f"\nГотово! Все файлы в директории '{BASE_DIR}/'")
    print("Теперь можно запустить краулер:")
    print("  python crawler.py --root storage/ --output output/index.csv --db output/fti.db")


if __name__ == "__main__":
    main()